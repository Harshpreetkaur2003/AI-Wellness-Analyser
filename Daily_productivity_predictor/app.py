# app.py

import streamlit as st
import pandas as pd
import pickle
from docx import Document
import io
import random
import numpy as np
import plotly.graph_objects as go
import os

# ---------------- PAGE CONFIG ----------------
st.set_page_config(
    page_title="AI Wellness & Performance Analyzer",
    page_icon="🧠",
    layout="wide"
)

# ---------------- DARK CINEMATIC BACKGROUND ----------------
st.markdown(
    """
    <style>
    .stApp {
        background: url('https://images.unsplash.com/photo-1613478209690-476a50e7c2d6?auto=format&fit=crop&w=1950&q=80') no-repeat center center fixed;
        background-size: cover;
        filter: brightness(0.3);
        color: white;
        font-family: 'Segoe UI', sans-serif;
    }
    h1, h2, h3 {
        color: #00F5FF;
        font-weight: 700;
        text-shadow: 2px 2px 8px #000000;
    }
    .glass {
        background: rgba(0,0,0,0.75);
        padding: 20px;
        border-radius: 20px;
        backdrop-filter: blur(20px);
        margin-bottom: 20px;
        box-shadow: 0 0 15px #000000;
    }
    .metric-box {
        background: rgba(0,0,0,0.75);
        border-radius: 15px;
        padding: 15px;
        text-align: center;
        box-shadow: 0 0 10px #111111;
    }
    </style>
    """,
    unsafe_allow_html=True
)

# ---------------- SECRET KEY LOGIN ----------------
VALID_KEYS = ["UnlockWellnessAI", "AIInsight2026", "SmartLifeKey"]

secret_key = st.text_input("Enter Secret Key to Access the Analyzer", type="password")
if st.button("Login"):
    if secret_key in VALID_KEYS:
        st.success("✅ Access Granted! Welcome to AI Wellness & Performance Analyzer")
    else:
        st.error("❌ Invalid Secret Key")
        st.stop()
else:
    st.stop()

# ---------------- APP TITLE ----------------
st.markdown('<div class="glass"><h1>🧠 AI Wellness & Performance Analyzer</h1></div>', unsafe_allow_html=True)
st.markdown('<div class="glass"><h3>Smart Lifestyle • Stress Prediction • Fitness • Career Blueprint</h3></div>', unsafe_allow_html=True)

# ---------------- LOAD MODEL ----------------
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
MODEL_PATH = os.path.join(BASE_DIR, "models", "productivity_model.pkl")
LE_PATH = os.path.join(BASE_DIR, "models", "label_encoder.pkl")

with open(MODEL_PATH, "rb") as f:
    model = pickle.load(f)

with open(LE_PATH, "rb") as f:
    le = pickle.load(f)

# ---------------- USER INPUT ----------------
st.header("📋 Enter Your Details")
name = st.text_input("Your Name")

col1, col2 = st.columns(2)
with col1:
    study_hours = st.number_input("Study Hours (per day)", 0, 12, 4)
    sleep_hours = st.number_input("Sleep Hours", 0, 12, 7)
    physical_activity = st.number_input("Workout Hours", 0, 4, 1)
    social_hours = st.number_input("Social Hours", 0, 6, 2)
    gpa = st.number_input("GPA", 0.0, 10.0, 8.0, 0.1)
    weight = st.number_input("Weight (kg)", 30, 200, 65)
    height = st.number_input("Height (cm)", 100, 220, 165)

with col2:
    career_stress = st.slider("Career Stress (1-10)", 1, 10, 5)
    motivation = st.slider("Motivation Level (1-10)", 1, 10, 7)
    water = st.number_input("Water Intake (liters)", 0.0, 5.0, 2.0, 0.1)
    food_type = st.selectbox("Diet Preference", ["Vegetarian", "Non-Vegetarian", "Vegan"])
    workout_type = st.selectbox("Workout Preference",
        ["Gym Training", "Home Workout", "Yoga Only", "Cardio Focus", "Mixed Routine"])

st.header("🎯 Career Direction")
career_domain = st.selectbox(
    "Select Career Domain",
    ["Management", "IT & Data", "Government Exams", "Creative Field", "Entrepreneurship"]
)
career_niche = st.text_input("Specific Niche (Example: Data Science, MBA Finance, UPSC, UI/UX)")

st.markdown("---")

# ---------------- GENERATE REPORT ----------------
if st.button("Generate Full AI Wellness Report"):
    st.balloons()

    # ML Prediction
    input_df = pd.DataFrame({
        "Study_Hours_Per_Day": [study_hours],
        "Extracurricular_Hours_Per_Day": [1],
        "Sleep_Hours_Per_Day": [sleep_hours],
        "Social_Hours_Per_Day": [social_hours],
        "Physical_Activity_Hours_Per_Day": [physical_activity],
        "GPA": [gpa]
    })
    prediction = model.predict(input_df)
    stress_level = le.inverse_transform(prediction)[0]

    # KPI Metrics
    bmi = weight / ((height / 100) ** 2)
    productivity_score = max(0, min(int((sleep_hours * 10) + (physical_activity * 15) + (motivation * 5) - (career_stress * 4)), 100))
    health_score = max(0, min(int((100 - abs(22 - bmi) * 5) + (physical_activity * 10) + (water * 5)), 100))

    st.markdown(f'<div class="glass"><h2>💬 Dear {name}, Here Is Your Detailed AI Analysis</h2></div>', unsafe_allow_html=True)
    colA, colB, colC = st.columns(3)
    colA.metric("Stress Level", stress_level)
    colB.metric("Productivity Score", f"{productivity_score}/100")
    colC.metric("Health Score", f"{health_score}/100")

    st.progress(productivity_score / 100)

    # ---------------- TABS ----------------
    tab1, tab2, tab3, tab4 = st.tabs(["📊 Analytics","🥗 Diet & Workout","🧠 Consultant Report","🎯 Career Blueprint"])

    # ---------------- TAB 1: Analytics ----------------
    with tab1:
        st.subheader("📈 Lifestyle & Performance Overview")
        factors = ["Sleep","Workout","Motivation","Stress Impact"]
        values = [sleep_hours*10, physical_activity*15, motivation*5, -career_stress*4]
        fig_combo = go.Figure()
        fig_combo.add_trace(go.Bar(x=factors, y=values, name="Impact Contribution", marker_color="#00F5FF"))
        fig_combo.add_trace(go.Scatter(x=factors, y=np.cumsum(values), mode="lines+markers", name="Cumulative Effect", line=dict(color="#FF2E63")))
        fig_combo.update_layout(template="plotly_dark", height=400, title="Factor Contribution & Cumulative Curve")
        st.plotly_chart(fig_combo, use_container_width=True)

        st.subheader("⚖ Daily Life Balance")
        balance_data = {"Study": study_hours,"Sleep": sleep_hours,"Workout": physical_activity,"Social": social_hours}
        fig_donut = go.Figure(data=[go.Pie(labels=list(balance_data.keys()), values=list(balance_data.values()), hole=0.6)])
        fig_donut.update_layout(template="plotly_dark", height=400, title="Daily Balance Ratio")
        st.plotly_chart(fig_donut, use_container_width=True)

        st.subheader("🧠 Predictive 3D Performance Surface")
        x = np.linspace(4,10,20)
        y = np.linspace(1,6,20)
        X,Y = np.meshgrid(x,y)
        Z = (X*8)+(Y*12)
        fig_surface = go.Figure(data=[go.Surface(z=Z,x=X,y=Y)])
        fig_surface.update_layout(template="plotly_dark", height=500, scene=dict(xaxis_title="Sleep Hours", yaxis_title="Workout Hours", zaxis_title="Performance Potential"), title="Lifestyle-Performance Simulation")
        st.plotly_chart(fig_surface,use_container_width=True)

    # ---------------- TAB 2: Diet & Workout ----------------
    with tab2:
        st.subheader("🥗 Personalized Nutrition & Meal Plan")
        if food_type == "Vegetarian":
            st.write("Breakfast: Oats + Milk + Almonds (Complex carbs + Protein)")
            st.write("Lunch: Dal + Brown Rice + Paneer (High protein vegetarian)")
            st.write("Snack: Fruits + Nuts")
            st.write("Dinner: Light roti + vegetables")
        elif food_type == "Vegan":
            st.write("Breakfast: Peanut Butter Smoothie (Plant protein)")
            st.write("Lunch: Quinoa + Chickpeas (Complete amino acids)")
            st.write("Snack: Seeds Mix")
            st.write("Dinner: Tofu + Vegetables")
        else:
            st.write("Breakfast: Eggs + Whole wheat toast (Protein + Carbs)")
            st.write("Lunch: Grilled Chicken + Rice (Muscle repair)")
            st.write("Snack: Yogurt")
            st.write("Dinner: Fish + Salad (Omega-3)")

        st.subheader("🏋 Structured Weekly Workout Plan")
        if workout_type == "Gym Training":
            st.write("Mon: Chest + Triceps\nTue: Back + Biceps\nWed: Legs\nThu: Shoulders\nFri: Core + HIIT")
        elif workout_type == "Home Workout":
            st.write("Pushups 3x15\nSquats 3x20\nPlank 3x60 sec\nJump Rope 10 min")
        elif workout_type == "Yoga Only":
            st.write("Surya Namaskar 10 rounds\nPranayama 15 min\nMeditation 20 min")
        elif workout_type == "Cardio Focus":
            st.write("Running 30 min\nCycling 20 min\nHIIT 15 min")
        else:
            st.write("Strength 3 days\nCardio 2 days\nYoga 1 day")

    # ---------------- TAB 3: Consultant Report ----------------
    with tab3:
        st.subheader("🧠 Detailed Consultant Analysis")
        st.write(f"Dear {name}, based on your lifestyle and inputs:")

        st.markdown("**Stress Analysis:**")
        if career_stress > 7:
            st.write("High stress! Incorporate relaxation & recovery cycles.")
        elif career_stress > 4:
            st.write("Moderate stress; manageable with proper planning.")
        else:
            st.write("Healthy stress zone; keep balanced routines.")

        st.markdown("**Productivity Deep Dive:**")
        st.write(f"- Sleep Contribution: {sleep_hours*10}\n- Workout Contribution: {physical_activity*15}\n- Motivation Contribution: {motivation*5}\n- Stress Deduction: {-career_stress*4}")

        st.markdown("**Nutrition Advice:**")
        st.write(f"{food_type} diet optimized for mental clarity, energy, and muscle recovery.")

        st.markdown("**Workout Guidance:**")
        st.write(f"{workout_type} routine structured for max performance and recovery.")

    # ---------------- TAB 4: Career Blueprint ----------------
    with tab4:
        st.subheader("🎯 Career Path & Skill Guide")
        st.write(f"Domain: {career_domain}\nSpecialization: {career_niche}")
        career_skills = {
            "IT & Data": ["Data Analysis", "Machine Learning", "AI Development", "Cloud Computing"],
            "Management": ["Leadership", "Strategic Planning", "Finance", "Project Management"],
            "Government Exams": ["General Knowledge", "Reasoning", "Current Affairs", "Subject Specialization"],
            "Creative Field": ["Design", "UI/UX", "Content Creation", "Brand Strategy"],
            "Entrepreneurship": ["Business Planning", "Marketing", "Financial Management", "Networking"]
        }
        st.write("**Core Skills Required:**")
        skills = career_skills.get(career_domain, [])
        for skill in skills:
            st.write(f"- {skill}")

        st.write("**Recommended Books & Resources:**")
        if career_domain == "IT & Data":
            st.write("- Python for Data Analysis (Wes McKinney)")
            st.write("- Hands-On Machine Learning with Scikit-Learn, Keras & TensorFlow (Aurelien Geron)")
        elif career_domain == "Management":
            st.write("- The Personal MBA (Josh Kaufman)")
            st.write("- Principles of Management (Harold Koontz)")
        elif career_domain == "Government Exams":
            st.write("- Lucent's General Knowledge")
            st.write("- Fast Track Objective Arithmetic (R.S. Aggarwal)")
        elif career_domain == "Creative Field":
            st.write("- Don't Make Me Think (Steve Krug)")
            st.write("- The Design of Everyday Things (Don Norman)")
        elif career_domain == "Entrepreneurship":
            st.write("- Zero to One (Peter Thiel)")
            st.write("- The Lean Startup (Eric Ries)")

    # ---------------- REPORT DOWNLOAD ----------------
    doc = Document()
    doc.add_heading("AI Wellness & Career Report", 0)
    doc.add_paragraph(f"Name: {name}")
    doc.add_paragraph(f"Stress Level: {stress_level}")
    doc.add_paragraph(f"Productivity Score: {productivity_score}")
    doc.add_paragraph(f"Health Score: {health_score}")
    doc.add_paragraph(f"Career Domain: {career_domain}")
    doc.add_paragraph(f"Career Niche: {career_niche}")

    buffer = io.BytesIO()
    doc.save(buffer)
    st.download_button(
        "⬇️ Download Full Professional Report",
        data=buffer.getvalue(),
        file_name="AI_Wellness_Report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
